# technical/management/commands/generate_club_report.py
"""
Django Management Command for CLUB Substation Report Generation
Usage: python manage.py generate_club_report
"""

from django.core.management.base import BaseCommand, CommandError
from django.db.models import Avg, Count, Max, Min, Sum, F, Q
from django.utils import timezone
from datetime import datetime, timedelta, date
from dateutil.relativedelta import relativedelta
import json
import os

from common.models import InjectionSubstation, Feeder
from technical.models import (
    HourlyLoad, FeederInterruption, DailyHoursOfSupply, 
    FeederEnergyDaily, FeederEnergyMonthly, calculate_interruption_metrics
)
from commercial.models import Customer
from analytics.models import MonthlyTechnicalSummary, DailyTechnicalSummary

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DateTimeEncoder(json.JSONEncoder):
    """Custom JSON encoder to handle date and datetime objects"""
    def default(self, obj):
        if isinstance(obj, (date, datetime)):
            return obj.isoformat()
        return super().default(obj)


class Command(BaseCommand):
    help = 'Generate comprehensive technical report for CLUB injection substation'

    def add_arguments(self, parser):
        parser.add_argument(
            '--output-dir',
            type=str,
            default='reports',
            help='Directory to save the report (default: reports)'
        )
        parser.add_argument(
            '--format',
            choices=['json', 'docx', 'summary', 'all'],
            default='all',
            help='Output format: json, docx, summary to console, or all (default: all)'
        )
        parser.add_argument(
            '--substation-id',
            type=str,
            default='0052cf6e-15ef-4de2-9c05-5895875ef791',
            help='Substation ID to generate report for (default: CLUB substation)'
        )
        parser.add_argument(
            '--report-date',
            type=str,
            default='2025-10-06',
            help='Report end date in YYYY-MM-DD format (default: 2025-10-06)'
        )

    def handle(self, *args, **options):
        self.stdout.write(
            self.style.SUCCESS('Starting CLUB Substation Report Generation...')
        )
        
        # Check if docx is requested but not available
        if options['format'] in ['docx', 'all'] and not DOCX_AVAILABLE:
            self.stdout.write(
                self.style.WARNING('python-docx not installed. Install with: pip install python-docx')
            )
            if options['format'] == 'docx':
                raise CommandError('python-docx is required for DOCX format. Install with: pip install python-docx')
            else:
                self.stdout.write(self.style.WARNING('Skipping DOCX generation...'))
        
        try:
            generator = CLUBSubstationReportGenerator(
                substation_id=options['substation_id'],
                report_date=options['report_date'],
                stdout=self.stdout,
                style=self.style
            )
            
            report = generator.generate_full_report()
            
            output_dir = options['output_dir']
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Handle JSON output
            if options['format'] in ['json', 'all']:
                filename = f"club_substation_report_{timestamp}.json"
                filepath = os.path.join(output_dir, filename)
                
                with open(filepath, 'w', encoding='utf-8') as f:
                    json.dump(report, f, indent=2, ensure_ascii=False, cls=DateTimeEncoder)
                
                self.stdout.write(
                    self.style.SUCCESS(f'JSON report saved to: {filepath}')
                )
            
            # Handle DOCX output
            if options['format'] in ['docx', 'all'] and DOCX_AVAILABLE:
                filename = f"club_substation_report_{timestamp}.docx"
                filepath = os.path.join(output_dir, filename)
                
                self.generate_docx_report(report, filepath)
                
                self.stdout.write(
                    self.style.SUCCESS(f'DOCX report saved to: {filepath}')
                )
            
            # Handle console summary
            if options['format'] in ['summary', 'all']:
                self.print_executive_summary(report)
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            raise CommandError(f'Report generation failed: {str(e)}')

    def generate_docx_report(self, report, filepath):
        """Generate a formatted Word document from the report data"""
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Title
        title = doc.add_heading('KEDCO CLUB SUBSTATION', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_heading('Raven Performance Monitoring Report', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary Section
        summary = report['executive_summary']
        deployment_info = summary['deployment_info']
        
        doc.add_heading('Executive Summary', 1)
        
        # Deployment Info
        p = doc.add_paragraph()
        p.add_run('Report Date: ').bold = True
        p.add_run(summary['report_date'])
        
        p = doc.add_paragraph()
        p.add_run('Deployment Status: ').bold = True
        p.add_run(deployment_info['status'])
        
        p = doc.add_paragraph()
        p.add_run('Operational Since: ').bold = True
        p.add_run(f"{deployment_info['deployment_date']} ({deployment_info['days_operational']} days)")
        
        p = doc.add_paragraph()
        p.add_run('Data Coverage: ').bold = True
        p.add_run(deployment_info['data_coverage'])
        
        p = doc.add_paragraph()
        p.add_run('Historical Baseline: ').bold = True
        p.add_run(deployment_info['historical_baseline'])
        
        # Infrastructure Overview
        doc.add_heading('Infrastructure Overview', 2)
        infra = summary['infrastructure_summary']
        voltage = infra['voltage_breakdown']
        
        infra_table = doc.add_table(rows=4, cols=2)
        infra_table.style = 'Light Grid Accent 1'
        
        infra_data = [
            ('Total Feeders', f"{infra['feeders']} ({voltage['11kv_feeders']} × 11kV, {voltage['33kv_feeders']} × 33kV)"),
            ('Distribution Transformers', str(infra['transformers'])),
            ('Customers Served', f"{infra['customers_served']:,}"),
            ('Substation', summary['substation'])
        ]
        
        for i, (label, value) in enumerate(infra_data):
            row_cells = infra_table.rows[i].cells
            row_cells[0].text = label
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].text = value
        
        # Key Performance Indicators
        doc.add_heading('Key Performance Indicators', 2)
        kpis = summary['key_performance_indicators']
        
        p = doc.add_paragraph()
        p.add_run(f"Period: {kpis['period']}").bold = True
        
        kpi_table = doc.add_table(rows=7, cols=2)
        kpi_table.style = 'Light List Accent 1'
        
        kpi_data = [
            ('Metric', 'Value'),
            ('Average Supply Hours', f"{kpis['average_supply_hours']:.1f} hrs/day"),
            ('System Availability', f"{kpis['system_availability']:.1f}%"),
            ('Total Interruptions', str(kpis['total_interruptions'])),
            ('Daily Interruption Rate', f"{kpis['daily_interruption_rate']:.1f} per day"),
            ('Peak Load', f"{kpis['peak_load_mw']:.1f} MW"),
            ('Energy Delivered', f"{kpis['energy_delivered_mwh']:.1f} MWh")
        ]
        
        for i, (metric, value) in enumerate(kpi_data):
            row_cells = kpi_table.rows[i].cells
            row_cells[0].text = metric
            row_cells[1].text = value
            if i == 0:  # Header row
                for cell in kpi_table.rows[i].cells:
                    cell.paragraphs[0].runs[0].bold = True
        
        # Deployment Performance
        if summary.get('deployment_performance'):
            doc.add_heading('Deployment Performance', 2)
            for item in summary['deployment_performance']:
                doc.add_paragraph(item, style='List Bullet')
        
        # Performance Highlights
        if summary.get('highlights'):
            doc.add_heading('Performance Highlights', 2)
            for item in summary['highlights']:
                p = doc.add_paragraph(item, style='List Bullet')
                p.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
        
        # Comparative Analysis
        if summary.get('comparative_insights'):
            doc.add_heading('Comparative Analysis (vs 2024 Baseline)', 2)
            for item in summary['comparative_insights']:
                doc.add_paragraph(item, style='List Bullet')
        
        # Areas of Concern
        if summary.get('concerns'):
            doc.add_heading('Areas of Concern', 2)
            for item in summary['concerns']:
                p = doc.add_paragraph(item, style='List Bullet')
                p.runs[0].font.color.rgb = RGBColor(255, 140, 0)  # Orange
        
        # Key Recommendations
        if summary.get('recommendations'):
            doc.add_heading('Key Recommendations', 2)
            for item in summary['recommendations']:
                doc.add_paragraph(item, style='List Bullet')
        
        # Detailed Comparisons
        doc.add_page_break()
        doc.add_heading('Detailed Performance Comparisons', 1)
        
        comparisons = report['performance_metrics'].get('key_comparisons', {})
        
        # 2024 Baseline Comparison
        baseline_comp = comparisons.get('deployment_vs_2024_baseline', {})
        if baseline_comp:
            doc.add_heading('Deployment Period vs 2024 Baseline', 2)
            doc.add_paragraph(baseline_comp.get('description', ''))
            
            comp_table = doc.add_table(rows=5, cols=4)
            comp_table.style = 'Medium Grid 3 Accent 1'
            
            # Headers
            headers = comp_table.rows[0].cells
            headers[0].text = 'Metric'
            headers[1].text = 'Current (2025)'
            headers[2].text = 'Baseline (2024)'
            headers[3].text = 'Change (%)'
            for cell in headers:
                cell.paragraphs[0].runs[0].bold = True
            
            # Supply Hours
            row = comp_table.rows[1].cells
            row[0].text = 'Supply Hours'
            row[1].text = f"{baseline_comp['supply_hours']['current']:.1f} hrs"
            row[2].text = f"{baseline_comp['supply_hours']['baseline_2024']:.1f} hrs"
            row[3].text = f"{baseline_comp['supply_hours']['improvement_percent']:+.1f}%"
            
            # Availability
            row = comp_table.rows[2].cells
            row[0].text = 'System Availability'
            row[1].text = f"{baseline_comp['availability']['current']:.1f}%"
            row[2].text = f"{baseline_comp['availability']['baseline_2024']:.1f}%"
            row[3].text = f"{baseline_comp['availability']['improvement_percent']:+.1f}%"
            
            # Interruptions
            row = comp_table.rows[3].cells
            row[0].text = 'Daily Interruptions'
            row[1].text = f"{baseline_comp['interruption_frequency']['current_daily_avg']:.1f}"
            row[2].text = f"{baseline_comp['interruption_frequency']['baseline_daily_avg']:.1f}"
            row[3].text = f"{baseline_comp['interruption_frequency']['change_percent']:+.1f}%"
            
            # Peak Load
            row = comp_table.rows[4].cells
            row[0].text = 'Peak Load'
            row[1].text = f"{baseline_comp['peak_load']['current']:.1f} MW"
            row[2].text = f"{baseline_comp['peak_load']['baseline_2024']:.1f} MW"
            row[3].text = f"{baseline_comp['peak_load']['change_percent']:+.1f}%"
        
        # Month-to-Month Progress
        monthly_comp = comparisons.get('august_to_september_2025', {})
        if monthly_comp:
            doc.add_heading('Month-to-Month Progress (2025)', 2)
            doc.add_paragraph(monthly_comp.get('description', ''))
            
            mom_table = doc.add_table(rows=3, cols=4)
            mom_table.style = 'Medium Grid 3 Accent 1'
            
            # Headers
            headers = mom_table.rows[0].cells
            headers[0].text = 'Metric'
            headers[1].text = 'Previous Month'
            headers[2].text = 'Current Month'
            headers[3].text = 'Change (%)'
            for cell in headers:
                cell.paragraphs[0].runs[0].bold = True
            
            # Supply Hours
            row = mom_table.rows[1].cells
            row[0].text = 'Supply Hours'
            row[1].text = f"{monthly_comp['supply_hours']['august']:.1f} hrs"
            row[2].text = f"{monthly_comp['supply_hours']['september']:.1f} hrs"
            row[3].text = f"{monthly_comp['supply_hours']['change_percent']:+.1f}%"
            
            # Interruptions
            row = mom_table.rows[2].cells
            row[0].text = 'Daily Interruptions'
            row[1].text = f"{monthly_comp['interruptions']['august_daily_avg']:.1f}"
            row[2].text = f"{monthly_comp['interruptions']['september_daily_avg']:.1f}"
            row[3].text = f"{monthly_comp['interruptions']['change_percent']:+.1f}%"
        
        # Infrastructure Details
        doc.add_page_break()
        doc.add_heading('Detailed Infrastructure Breakdown', 1)
        
        infra_details = report['infrastructure_overview']
        
        # Band Distribution
        doc.add_heading('Feeder Distribution by Service Band', 2)
        band_dist = infra_details['band_distribution']
        for band, count in band_dist.items():
            doc.add_paragraph(f"{band}: {count} feeders", style='List Bullet')
        
        # Feeder Details Table
        doc.add_heading('Complete Feeder List', 2)
        feeders = infra_details['feeder_details']
        
        if feeders:
            feeder_table = doc.add_table(rows=len(feeders) + 1, cols=6)
            feeder_table.style = 'Light Grid Accent 1'
            
            # Headers
            headers = feeder_table.rows[0].cells
            headers[0].text = 'Feeder Name'
            headers[1].text = 'Voltage'
            headers[2].text = 'Band'
            headers[3].text = 'District'
            headers[4].text = 'Transformers'
            headers[5].text = 'Customers'
            for cell in headers:
                cell.paragraphs[0].runs[0].bold = True
            
            # Data rows
            for i, feeder in enumerate(feeders, start=1):
                row = feeder_table.rows[i].cells
                row[0].text = feeder['name']
                row[1].text = feeder['voltage_level']
                row[2].text = feeder['band']
                row[3].text = feeder['business_district']
                row[4].text = str(feeder['transformer_count'])
                row[5].text = str(feeder['customer_count'])
        
        # Performance Metrics by Period
        doc.add_page_break()
        doc.add_heading('Performance Metrics by Period', 1)
        
        metrics_by_period = report['performance_metrics']['metrics_by_period']
        
        # Key periods to report
        key_periods = [
            ('since_deployment', 'Since Deployment (Aug 9 - Oct 6, 2025)'),
            ('october_2025', 'October 2025'),
            ('september_2025', 'September 2025'),
            ('last_7_days', 'Last 7 Days'),
            ('last_14_days', 'Last 14 Days'),
            ('last_30_days', 'Last 30 Days'),
            ('august_2024', 'August 2024 (Baseline)'),
            ('september_2024', 'September 2024 (Baseline)'),
            ('october_2024_partial', 'October 2024 (Baseline)'),
            ('full_year_2024', 'Full Year 2024'),
            ('q3_2024', 'Q3 2024 (Baseline)')
        ]
        
        for period_key, period_label in key_periods:
            if period_key in metrics_by_period:
                period_data = metrics_by_period[period_key]
                
                doc.add_heading(period_label, 2)
                
                period_info = period_data['period']
                doc.add_paragraph(f"Period: {period_info['start_date']} to {period_info['end_date']} ({period_info['days']} days)")
                
                # Metrics table
                metrics_table = doc.add_table(rows=8, cols=2)
                metrics_table.style = 'Light List Accent 1'
                
                supply = period_data['supply_hours']
                load = period_data['load_metrics']
                energy = period_data['energy_delivered']
                interruptions = period_data['interruptions']
                reliability = period_data['reliability']
                
                metrics_rows = [
                    ('Average Supply Hours', f"{supply['average']:.1f} hrs/day (Max: {supply['maximum']:.1f}, Min: {supply['minimum']:.1f})"),
                    ('Peak Load', f"{load['peak_load_mw']:.1f} MW"),
                    ('Average Load', f"{load['average_load_mw']:.1f} MW"),
                    ('Energy Delivered', f"{energy['total_mwh']:.1f} MWh (Avg Daily: {energy.get('average_daily_mwh', 0):.1f} MWh)"),
                    ('Total Interruptions', str(interruptions['total_interruptions'])),
                    ('Fault Count', str(interruptions['fault_count'])),
                    ('Load Shedding Events', str(interruptions['load_shedding_count'])),
                    ('System Availability', f"{reliability['availability_percentage']:.1f}%")
                ]
                
                for i, (label, value) in enumerate(metrics_rows):
                    row = metrics_table.rows[i].cells
                    row[0].text = label
                    row[0].paragraphs[0].runs[0].bold = True
                    row[1].text = value
                
                # Reliability Indices
                doc.add_paragraph('Reliability Indices:', style='Heading 3')
                reliability_table = doc.add_table(rows=5, cols=2)
                reliability_table.style = 'Light Grid Accent 1'
                
                reliability_rows = [
                    ('Index', 'Value'),
                    ('SAIFI (interruptions/customer)', f"{reliability['saifi']:.4f}"),
                    ('SAIDI (hours/customer)', f"{reliability['saidi']:.4f}"),
                    ('CAIDI (hours/interruption)', f"{reliability['caidi']:.4f}"),
                    ('Customer Base', f"{reliability['total_customer_base']:,}")
                ]
                
                for i, (label, value) in enumerate(reliability_rows):
                    row = reliability_table.rows[i].cells
                    row[0].text = label
                    row[1].text = value
                    if i == 0:
                        row[0].paragraphs[0].runs[0].bold = True
                        row[1].paragraphs[0].runs[0].bold = True
                
                # Interruption Type Breakdown
                if 'breakdown_by_type' in interruptions and interruptions['breakdown_by_type']:
                    doc.add_paragraph('Interruption Breakdown by Type:', style='Heading 3')
                    breakdown = interruptions['breakdown_by_type']
                    
                    # Sort by count descending
                    sorted_types = sorted(breakdown.items(), key=lambda x: x[1]['count'], reverse=True)
                    
                    if sorted_types:
                        breakdown_table = doc.add_table(rows=len(sorted_types) + 1, cols=4)
                        breakdown_table.style = 'Light Grid Accent 1'
                        
                        # Headers
                        headers = breakdown_table.rows[0].cells
                        headers[0].text = 'Type'
                        headers[1].text = 'Count'
                        headers[2].text = 'Total Hours'
                        headers[3].text = 'Avg Duration'
                        for cell in headers:
                            cell.paragraphs[0].runs[0].bold = True
                        
                        # Data rows
                        for i, (int_type, int_data) in enumerate(sorted_types, start=1):
                            row = breakdown_table.rows[i].cells
                            row[0].text = int_type
                            row[1].text = str(int_data['count'])
                            row[2].text = f"{int_data['duration']:.2f}"
                            row[3].text = f"{int_data['avg_duration']:.2f}"
                
                # Peak Load by Feeder
                if 'peak_loads_by_feeder' in load and load['peak_loads_by_feeder']:
                    doc.add_paragraph('Peak Load by Feeder (Top 10):', style='Heading 3')
                    peak_table = doc.add_table(rows=len(load['peak_loads_by_feeder']) + 1, cols=3)
                    peak_table.style = 'Light List Accent 1'
                    
                    # Headers
                    headers = peak_table.rows[0].cells
                    headers[0].text = 'Feeder'
                    headers[1].text = 'Voltage'
                    headers[2].text = 'Peak Load (MW)'
                    for cell in headers:
                        cell.paragraphs[0].runs[0].bold = True
                    
                    # Data rows
                    for i, feeder_data in enumerate(load['peak_loads_by_feeder'], start=1):
                        row = peak_table.rows[i].cells
                        row[0].text = feeder_data['feeder']
                        row[1].text = feeder_data['voltage_level']
                        row[2].text = f"{feeder_data['peak_load_mw']:.2f}"
                
                # Energy by feeder
                if 'energy_by_feeder' in energy and energy['energy_by_feeder']:
                    doc.add_paragraph('Top Energy-Delivering Feeders:', style='Heading 3')
                    energy_table = doc.add_table(rows=len(energy['energy_by_feeder']) + 1, cols=2)
                    energy_table.style = 'Light List Accent 1'
                    
                    # Headers
                    headers = energy_table.rows[0].cells
                    headers[0].text = 'Feeder'
                    headers[1].text = 'Energy (MWh)'
                    for cell in headers:
                        cell.paragraphs[0].runs[0].bold = True
                    
                    # Data rows
                    for i, feeder_energy in enumerate(energy['energy_by_feeder'], start=1):
                        row = energy_table.rows[i].cells
                        row[0].text = feeder_energy['feeder_name']
                        row[1].text = f"{feeder_energy['energy_mwh']:.1f}"
                
                # Calculation metadata
                doc.add_paragraph(f"Energy Calculation: {energy.get('calculation_method', 'N/A')}", style='Body Text')
                doc.add_paragraph(f"Data Points Used: {energy.get('total_load_records_used', 0):,} load records", style='Body Text')
        
        # Feeder-Specific Interruption Analysis
        doc.add_page_break()
        doc.add_heading('Feeder-Specific Interruption Analysis', 1)
        
        # Get the since_deployment metrics for detailed feeder breakdown
        deployment_data = metrics_by_period.get('since_deployment', {})
        if deployment_data and 'interruptions' in deployment_data:
            feeder_breakdown = deployment_data['interruptions'].get('feeder_breakdown', {})
            
            if feeder_breakdown:
                doc.add_paragraph('Interruption metrics for each feeder during the deployment period:')
                
                # Sort feeders by total interruptions descending
                sorted_feeders = sorted(feeder_breakdown.items(), 
                                      key=lambda x: x[1]['total_interruptions'], 
                                      reverse=True)
                
                for feeder_name, feeder_metrics in sorted_feeders:
                    doc.add_heading(feeder_name, 3)
                    
                    feeder_table = doc.add_table(rows=7, cols=2)
                    feeder_table.style = 'Light List Accent 1'
                    
                    feeder_rows = [
                        ('Total Interruptions', str(feeder_metrics['total_interruptions'])),
                        ('Total Duration (hours)', f"{feeder_metrics['total_duration_hours']:.2f}"),
                        ('Average Duration (hours)', f"{feeder_metrics['avg_duration_hours']:.2f}"),
                        ('Fault Count', str(feeder_metrics['fault_count'])),
                        ('Fault Hours', f"{feeder_metrics['fault_hours']:.2f}"),
                        ('Load Shedding Count', str(feeder_metrics['load_shedding_count'])),
                        ('Load Shedding Hours', f"{feeder_metrics['load_shedding_hours']:.2f}")
                    ]
                    
                    for i, (label, value) in enumerate(feeder_rows):
                        row = feeder_table.rows[i].cells
                        row[0].text = label
                        row[0].paragraphs[0].runs[0].bold = True
                        row[1].text = value
        
        # Daily Interruption Trend
        doc.add_page_break()
        doc.add_heading('Daily Interruption Trends', 1)
        
        if deployment_data and 'interruptions' in deployment_data:
            daily_trend = deployment_data['interruptions'].get('daily_trend', [])
            
            if daily_trend:
                doc.add_paragraph('Daily interruption counts throughout the deployment period:')
                
                # Create table for daily data (split into chunks if too long)
                chunk_size = 30  # 30 days per table
                for chunk_start in range(0, len(daily_trend), chunk_size):
                    chunk = daily_trend[chunk_start:chunk_start + chunk_size]
                    
                    trend_table = doc.add_table(rows=len(chunk) + 1, cols=2)
                    trend_table.style = 'Light Grid Accent 1'
                    
                    # Headers
                    headers = trend_table.rows[0].cells
                    headers[0].text = 'Date'
                    headers[1].text = 'Interruption Count'
                    for cell in headers:
                        cell.paragraphs[0].runs[0].bold = True
                    
                    # Data rows
                    for i, day_data in enumerate(chunk, start=1):
                        row = trend_table.rows[i].cells
                        row[0].text = day_data['date']
                        row[1].text = str(day_data['interruption_count'])
                    
                    # Add spacing between chunks
                    if chunk_start + chunk_size < len(daily_trend):
                        doc.add_paragraph('')  # Empty paragraph for spacing
        
        # Operational Insights
        doc.add_page_break()
        doc.add_heading('Operational Insights & Analysis', 1)
        
        insights = report['operational_insights']
        
        for insight_category, insight_list in insights.items():
            if insight_list:
                # Format category name
                category_name = insight_category.replace('_', ' ').title()
                doc.add_heading(category_name, 2)
                
                for item in insight_list:
                    doc.add_paragraph(item, style='List Bullet')
        
        # Metadata footer
        doc.add_page_break()
        doc.add_heading('Report Metadata', 1)
        
        metadata = report['metadata']
        
        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.style = 'Light Grid'
        
        meta_data = [
            ('Generated At', metadata['generated_at']),
            ('Generator Version', metadata['generator_version']),
            ('Substation ID', metadata['substation_id']),
            ('Data Sources', ', '.join(metadata['data_sources']))
        ]
        
        for i, (label, value) in enumerate(meta_data):
            row = meta_table.rows[i].cells
            row[0].text = label
            row[0].paragraphs[0].runs[0].bold = True
            row[1].text = value
        
        # Save document
        doc.save(filepath)

    def print_executive_summary(self, report):
        """Print formatted executive summary to console"""
        self.stdout.write("\n" + "=" * 80)
        self.stdout.write(
            self.style.HTTP_SUCCESS("KEDCO CLUB SUBSTATION - RAVEN DEPLOYMENT REPORT")
        )
        self.stdout.write("=" * 80)
        
        summary = report['executive_summary']
        deployment_info = summary['deployment_info']
        
        self.stdout.write(f"Substation: {summary['substation']}")
        self.stdout.write(f"Report Date: {summary['report_date']}")
        self.stdout.write(f"Deployment Status: {deployment_info['status']}")
        self.stdout.write(f"Operational Since: {deployment_info['deployment_date']} ({deployment_info['days_operational']} days)")
        self.stdout.write(f"Data Coverage: {deployment_info['data_coverage']}")
        self.stdout.write(f"Historical Context: {deployment_info['historical_baseline']}")
        
        self.stdout.write(f"\nINFRASTRUCTURE OVERVIEW:")
        infra = summary['infrastructure_summary']
        voltage = infra['voltage_breakdown']
        self.stdout.write(f"  • Total Feeders: {infra['feeders']} ({voltage['11kv_feeders']} x 11kV, {voltage['33kv_feeders']} x 33kV)")
        self.stdout.write(f"  • Distribution Transformers: {infra['transformers']}")
        self.stdout.write(f"  • Customers Served: {infra['customers_served']:,}")
        
        self.stdout.write(f"\nKEY PERFORMANCE INDICATORS ({summary['key_performance_indicators']['period']}):")
        kpis = summary['key_performance_indicators']
        self.stdout.write(f"  • Average Supply Hours: {kpis['average_supply_hours']:.1f} hrs/day")
        self.stdout.write(f"  • System Availability: {kpis['system_availability']:.1f}%")
        self.stdout.write(f"  • Total Interruptions: {kpis['total_interruptions']}")
        self.stdout.write(f"  • Daily Interruption Rate: {kpis['daily_interruption_rate']:.1f} per day")
        self.stdout.write(f"  • Peak Load: {kpis['peak_load_mw']:.1f} MW")
        self.stdout.write(f"  • Energy Delivered: {kpis['energy_delivered_mwh']:.1f} MWh")
        
        if summary.get('deployment_performance'):
            self.stdout.write(f"\nDEPLOYMENT PERFORMANCE:")
            for performance in summary['deployment_performance']:
                self.stdout.write(f"  • {performance}")
        
        if summary.get('highlights'):
            self.stdout.write(f"\nPERFORMANCE HIGHLIGHTS:")
            for highlight in summary['highlights']:
                self.stdout.write(
                    f"  • " + self.style.SUCCESS(highlight)
                )
        
        if summary.get('comparative_insights'):
            self.stdout.write(f"\nCOMPARATIVE ANALYSIS:")
            for insight in summary['comparative_insights']:
                self.stdout.write(f"  • {insight}")
        
        if summary.get('concerns'):
            self.stdout.write(f"\nAREAS OF CONCERN:")
            for concern in summary['concerns']:
                self.stdout.write(
                    f"  • " + self.style.WARNING(concern)
                )
        
        if summary.get('recommendations'):
            self.stdout.write(f"\nKEY RECOMMENDATIONS:")
            for rec in summary['recommendations']:
                self.stdout.write(f"  • {rec}")
        
        # Print detailed comparison data if available
        comparisons = report['performance_metrics'].get('key_comparisons', {})
        baseline_comp = comparisons.get('deployment_vs_2024_baseline', {})
        
        if baseline_comp:
            self.stdout.write(f"\nDETAILED 2024 BASELINE COMPARISON:")
            supply_change = baseline_comp['supply_hours']['improvement_percent']
            avail_change = baseline_comp['availability']['improvement_percent']
            int_change = baseline_comp['interruption_frequency']['change_percent']
            
            self.stdout.write(f"  Supply Hours: {baseline_comp['supply_hours']['current']:.1f}h vs {baseline_comp['supply_hours']['baseline_2024']:.1f}h (2024) = {supply_change:+.1f}%")
            self.stdout.write(f"  System Availability: {baseline_comp['availability']['current']:.1f}% vs {baseline_comp['availability']['baseline_2024']:.1f}% (2024) = {avail_change:+.1f}%")
            self.stdout.write(f"  Daily Interruptions: {baseline_comp['interruption_frequency']['current_daily_avg']:.1f} vs {baseline_comp['interruption_frequency']['baseline_daily_avg']:.1f} (2024) = {int_change:+.1f}%")
        
        monthly_comp = comparisons.get('august_to_september_2025', {})
        if monthly_comp:
            self.stdout.write(f"\nMONTH-TO-MONTH PROGRESS (2025):")
            prev_supply = monthly_comp['supply_hours']['august']
            curr_supply = monthly_comp['supply_hours']['september'] 
            supply_change = monthly_comp['supply_hours']['change_percent']
            self.stdout.write(f"  Supply Hours: Previous {prev_supply:.1f}h → Current {curr_supply:.1f}h ({supply_change:+.1f}%)")
        
        self.stdout.write("\n" + "=" * 80)


class CLUBSubstationReportGenerator:
    """Generate comprehensive technical report for CLUB substation"""
    
    def __init__(self, substation_id="0052cf6e-15ef-4de2-9c05-5895875ef791", report_date=None, stdout=None, style=None):
        self.substation_id = substation_id
        self.substation_slug = "KN-CLU"
        self.substation_name = "CLUB"
        self.stdout = stdout
        self.style = style
        
        # KEDCO deployment timeline
        self.deployment_start = datetime(2025, 8, 9).date()  # Raven deployment start
        
        # Parse report date
        if report_date:
            if isinstance(report_date, str):
                try:
                    self.current_date = datetime.strptime(report_date, '%Y-%m-%d').date()
                except ValueError:
                    raise ValueError(f"Invalid date format: {report_date}. Use YYYY-MM-DD")
            else:
                self.current_date = report_date
        else:
            self.current_date = datetime(2025, 10, 6).date()  # Default date
        
        self.historical_data_available = True  # Complete 2024 data available
        
        # Calculate deployment duration
        self.days_since_deployment = (self.current_date - self.deployment_start).days + 1
        
        # Get substation and feeders
        try:
            self.substation = InjectionSubstation.objects.get(id=self.substation_id)
            self.feeders = Feeder.objects.filter(substation=self.substation)
            self.feeder_ids = list(self.feeders.values_list('id', flat=True))
            
            if self.stdout:
                self.stdout.write(f"Found {len(self.feeder_ids)} feeders under {self.substation_name} substation")
                self.stdout.write(f"Deployment period: {self.deployment_start} to {self.current_date} ({self.days_since_deployment} days)")
        except InjectionSubstation.DoesNotExist:
            raise ValueError(f"Substation with ID {self.substation_id} not found")
    
    def get_date_ranges(self):
        """Generate date ranges based on actual deployment timeline and available data"""
        today = self.current_date
        deployment_start = self.deployment_start
        
        # 2025 ranges (post-deployment only)
        august_2025_start = max(deployment_start, datetime(2025, 8, 1).date())
        august_2025_end = datetime(2025, 8, 31).date()
        
        september_2025_start = datetime(2025, 9, 1).date()
        september_2025_end = datetime(2025, 9, 30).date()
        
        october_2025_start = datetime(2025, 10, 1).date()
        october_2025_end = min(today, datetime(2025, 10, 31).date())
        
        # 2024 ranges (complete historical data)
        august_2024 = (datetime(2024, 8, 1).date(), datetime(2024, 8, 31).date())
        september_2024 = (datetime(2024, 9, 1).date(), datetime(2024, 9, 30).date())
        october_2024 = (datetime(2024, 10, 1).date(), datetime(2024, 10, 6).date())  # Same period last year
        
        ranges = {
            # Current deployment period analysis
            'since_deployment': (deployment_start, today),
            'deployment_august_portion': (august_2025_start, august_2025_end),
            'september_2025': (september_2025_start, september_2025_end),
            'october_2025': (october_2025_start, october_2025_end),
            
            # Recent periods
            'last_7_days': (today - timedelta(days=6), today),
            'last_14_days': (today - timedelta(days=13), today),
            'last_30_days': (max(deployment_start, today - timedelta(days=29)), today),
            
            # Historical comparison periods (2024)
            'august_2024': august_2024,
            'september_2024': september_2024,
            'october_2024_partial': october_2024,  # Same period last year
            'august_october_2024': (august_2024[0], october_2024[1]),  # Full comparison period
            
            # Full 2024 for baseline
            'full_year_2024': (datetime(2024, 1, 1).date(), datetime(2024, 12, 31).date()),
            'q3_2024': (datetime(2024, 7, 1).date(), datetime(2024, 9, 30).date()),
        }
        
        # Add weekly breakdown of deployment period
        current_week_start = today - timedelta(days=today.weekday())
        if current_week_start >= deployment_start:
            ranges['current_week'] = (current_week_start, today)
        
        previous_week_start = current_week_start - timedelta(days=7)
        previous_week_end = current_week_start - timedelta(days=1)
        if previous_week_start >= deployment_start:
            ranges['previous_week'] = (previous_week_start, previous_week_end)
        
        return ranges
    
    def calculate_delta(self, current, previous):
        """Calculate percentage change between current and previous values"""
        if previous == 0:
            return 100 if current > 0 else 0
        return round(((current - previous) / previous) * 100, 2)
    
    def get_infrastructure_overview(self):
        """Get basic infrastructure information"""
        if self.stdout:
            self.stdout.write("Generating infrastructure overview...")
        
        feeder_details = []
        for feeder in self.feeders:
            # Get transformer count per feeder
            transformer_count = feeder.transformers.count()
            
            # Get customer count per feeder (if available)
            try:
                customer_count = Customer.objects.filter(
                    transformer__feeder=feeder
                ).count()
            except:
                customer_count = 0
            
            feeder_details.append({
                'name': feeder.name,
                'slug': feeder.slug,
                'voltage_level': feeder.voltage_level,
                'band': feeder.band.name if feeder.band else 'Unassigned',
                'business_district': feeder.business_district.name if feeder.business_district else 'Unassigned',
                'transformer_count': transformer_count,
                'customer_count': customer_count
            })
        
        total_transformers = sum(f['transformer_count'] for f in feeder_details)
        total_customers = sum(f['customer_count'] for f in feeder_details)
        
        return {
            'substation_name': self.substation_name,
            'substation_slug': self.substation_slug,
            'total_feeders': len(self.feeders),
            'total_transformers': total_transformers,
            'total_customers': total_customers,
            'feeder_breakdown': {
                '11kv_feeders': len([f for f in feeder_details if f['voltage_level'] == '11kv']),
                '33kv_feeders': len([f for f in feeder_details if f['voltage_level'] == '33kv'])
            },
            'band_distribution': self._get_band_distribution(feeder_details),
            'feeder_details': feeder_details
        }
    
    def _get_band_distribution(self, feeder_details):
        """Get distribution of feeders by band"""
        bands = {}
        for feeder in feeder_details:
            band = feeder['band']
            bands[band] = bands.get(band, 0) + 1
        return bands
    
    def get_performance_metrics(self):
        """Get comprehensive performance metrics for various time periods"""
        if self.stdout:
            self.stdout.write("Calculating performance metrics...")
        
        date_ranges = self.get_date_ranges()
        metrics = {}
        
        for period_name, (start_date, end_date) in date_ranges.items():
            if self.stdout:
                self.stdout.write(f"Processing {period_name}: {start_date} to {end_date}")
            
            period_metrics = self._calculate_period_metrics(start_date, end_date)
            metrics[period_name] = period_metrics
        
        # Calculate key comparisons
        comparisons = self._calculate_comparisons(metrics)
        
        return {
            'metrics_by_period': metrics,
            'key_comparisons': comparisons
        }
    
    def _calculate_period_metrics(self, start_date, end_date):
        """Calculate metrics for a specific period"""
        
        # Hours of Supply
        avg_supply = self._calculate_avg_supply_hours(start_date, end_date)
        
        # Load Metrics
        load_metrics = self._calculate_load_metrics(start_date, end_date)
        
        # Energy Delivered
        energy_delivered = self._calculate_energy_delivered(start_date, end_date)
        
        # Interruption Metrics
        interruption_metrics = self._calculate_interruption_metrics(start_date, end_date)
        
        # Reliability Indices
        reliability_indices = self._calculate_reliability_indices(
            start_date, end_date, avg_supply, interruption_metrics
        )
        
        return {
            'period': {
                'start_date': start_date.isoformat(),
                'end_date': end_date.isoformat(),
                'days': (end_date - start_date).days + 1
            },
            'supply_hours': avg_supply,
            'load_metrics': load_metrics,
            'energy_delivered': energy_delivered,
            'interruptions': interruption_metrics,
            'reliability': reliability_indices
        }
    
    def _calculate_avg_supply_hours(self, start_date, end_date):
        """Calculate average supply hours"""
        # Try DailyHoursOfSupply first
        daily_supply = DailyHoursOfSupply.objects.filter(
            feeder_id__in=self.feeder_ids,
            date__range=(start_date, end_date)
        ).aggregate(
            avg_hours=Avg('hours_supplied'),
            max_hours=Max('hours_supplied'),
            min_hours=Min('hours_supplied')
        )
        
        if daily_supply['avg_hours'] is not None:
            return {
                'average': round(float(daily_supply['avg_hours']), 2),
                'maximum': round(float(daily_supply['max_hours']), 2),
                'minimum': round(float(daily_supply['min_hours']), 2)
            }
        
        # Fallback to HourlyLoad calculation
        hourly_supply = HourlyLoad.objects.filter(
            feeder_id__in=self.feeder_ids,
            date__range=(start_date, end_date),
            load_mw__gt=0
        ).values('feeder', 'date').annotate(
            hours_count=Count('hour')
        ).aggregate(
            avg_hours=Avg('hours_count'),
            max_hours=Max('hours_count'),
            min_hours=Min('hours_count')
        )
        
        return {
            'average': round(float(hourly_supply['avg_hours'] or 0), 2),
            'maximum': round(float(hourly_supply['max_hours'] or 0), 2),
            'minimum': round(float(hourly_supply['min_hours'] or 0), 2)
        }
    
    def _calculate_load_metrics(self, start_date, end_date):
        """Calculate load-related metrics"""
        load_stats = HourlyLoad.objects.filter(
            feeder_id__in=self.feeder_ids,
            date__range=(start_date, end_date)
        ).aggregate(
            avg_load=Avg('load_mw'),
            max_load=Max('load_mw'),
            min_load=Min('load_mw'),
            total_records=Count('id')
        )
        
        # Peak load by feeder
        peak_loads_by_feeder = HourlyLoad.objects.filter(
            feeder_id__in=self.feeder_ids,
            date__range=(start_date, end_date)
        ).values(
            'feeder__name', 'feeder__voltage_level'
        ).annotate(
            peak_load=Max('load_mw')
        ).order_by('-peak_load')
        
        # Calculate Station Total Load per Hour
        # Group by date and hour, then sum load_mw across all feeders
        station_hourly_loads = HourlyLoad.objects.filter(
            feeder_id__in=self.feeder_ids,
            date__range=(start_date, end_date)
        ).values('date', 'hour').annotate(
            total_station_load=Sum('load_mw')
        ).order_by('date', 'hour')

        # Calculate Average Station Load
        # Average of the hourly station totals
        station_load_stats = station_hourly_loads.aggregate(
            avg_station_load=Avg('total_station_load'),
            max_station_load=Max('total_station_load'),
            min_station_load=Min('total_station_load')
        )

        return {
            'average_load_mw': round(float(load_stats['avg_load'] or 0), 2),
            'peak_load_mw': round(float(load_stats['max_load'] or 0), 2),
            'minimum_load_mw': round(float(load_stats['min_load'] or 0), 2),
            'total_load_records': load_stats['total_records'],
            
            # New Station-Level Metrics
            'average_station_load_mw': round(float(station_load_stats['avg_station_load'] or 0), 2),
            'peak_station_load_mw': round(float(station_load_stats['max_station_load'] or 0), 2),
            'min_station_load_mw': round(float(station_load_stats['min_station_load'] or 0), 2),
            'station_hourly_loads': [
                {
                    'date': item['date'].isoformat(),
                    'hour': item['hour'],
                    'total_load_mw': round(float(item['total_station_load'] or 0), 2)
                }
                for item in station_hourly_loads
            ],

            'peak_loads_by_feeder': [
                {
                    'feeder': item['feeder__name'],
                    'voltage_level': item['feeder__voltage_level'],
                    'peak_load_mw': round(float(item['peak_load'] or 0), 2)
                }
                for item in peak_loads_by_feeder[:10]  # Top 10
            ]
        }
    
    def _calculate_energy_delivered(self, start_date, end_date):
        """Calculate energy delivered from hourly load data (Energy = Power × Time)"""
        
        # Calculate energy from load data: Sum(MW × 1 hour) = MWh
        load_energy_data = HourlyLoad.objects.filter(
            feeder_id__in=self.feeder_ids,
            date__range=(start_date, end_date)
        ).aggregate(
            total_energy_mwh=Sum('load_mw'),  # Sum of MW readings = MWh
            total_records=Count('id')
        )
        
        total_energy = load_energy_data['total_energy_mwh'] or 0
        total_records = load_energy_data['total_records'] or 0
        
        # Calculate daily average
        days = (end_date - start_date).days + 1
        avg_daily_energy = total_energy / days if days > 0 else 0
        
        # Get energy breakdown by feeder
        energy_by_feeder = []
        for feeder in self.feeders[:10]:  # Top 10 feeders to avoid too much data
            feeder_energy = HourlyLoad.objects.filter(
                feeder=feeder,
                date__range=(start_date, end_date)
            ).aggregate(
                total=Sum('load_mw')
            )['total'] or 0
            
            if feeder_energy > 0:
                energy_by_feeder.append({
                    'feeder_name': feeder.name,
                    'energy_mwh': round(float(feeder_energy), 2)
                })
        
        # Sort by energy delivered
        energy_by_feeder.sort(key=lambda x: x['energy_mwh'], reverse=True)
        
        return {
            'total_mwh': round(float(total_energy), 2),
            'average_daily_mwh': round(float(avg_daily_energy), 2),
            'data_source': 'calculated_from_hourly_load_data',
            'calculation_method': 'Sum(MW × 1hour) = MWh',
            'total_load_records_used': total_records,
            'energy_by_feeder': energy_by_feeder[:5]  # Top 5 feeders
        }
    
    def _calculate_interruption_metrics(self, start_date, end_date):
        """Calculate detailed interruption metrics"""
        interruptions_qs = FeederInterruption.objects.filter(
            feeder_id__in=self.feeder_ids,
            occurred_at__date__range=(start_date, end_date)
        )
        
        # Use the utility function from models
        metrics = calculate_interruption_metrics(interruptions_qs)
        
        # Add feeder-specific breakdown
        feeder_breakdown = {}
        for feeder in self.feeders:
            feeder_interruptions = interruptions_qs.filter(feeder=feeder)
            feeder_metrics = calculate_interruption_metrics(feeder_interruptions)
            feeder_breakdown[feeder.name] = feeder_metrics
        
        # Add daily trend
        daily_trend = self._get_daily_interruption_trend(start_date, end_date)
        
        return {
            **metrics,
            'feeder_breakdown': feeder_breakdown,
            'daily_trend': daily_trend
        }
    
    def _get_daily_interruption_trend(self, start_date, end_date):
        """Get daily interruption counts for trend analysis"""
        daily_counts = FeederInterruption.objects.filter(
            feeder_id__in=self.feeder_ids,
            occurred_at__date__range=(start_date, end_date)
        ).extra(
            select={'day': 'DATE(occurred_at)'}
        ).values('day').annotate(
            count=Count('id')
        ).order_by('day')
        
        return [
            {
                'date': str(item['day']),  # Convert date to string for JSON serialization
                'interruption_count': item['count']
            }
            for item in daily_counts
        ]
    
    def _calculate_reliability_indices(self, start_date, end_date, supply_metrics, interruption_metrics):
        """Calculate standard reliability indices"""
        days_in_period = (end_date - start_date).days + 1
        
        # System Average Interruption Frequency Index (SAIFI)
        try:
            total_customers = sum(
                Customer.objects.filter(transformer__feeder=feeder).count() 
                for feeder in self.feeders
            )
        except:
            total_customers = 1  # Avoid division by zero
        
        if total_customers == 0:
            total_customers = 1
        
        saifi = interruption_metrics['total_interruptions'] / total_customers
        
        # System Average Interruption Duration Index (SAIDI)
        saidi = interruption_metrics['total_duration_hours'] / total_customers
        
        # Customer Average Interruption Duration Index (CAIDI)
        caidi = (interruption_metrics['total_duration_hours'] / interruption_metrics['total_interruptions'] 
                if interruption_metrics['total_interruptions'] > 0 else 0)
        
        # Availability percentage (based on supply hours)
        availability = (supply_metrics['average'] / 24) * 100 if supply_metrics['average'] else 0
        
        return {
            'saifi': round(saifi, 4),
            'saidi': round(saidi, 4),
            'caidi': round(caidi, 4),
            'availability_percentage': round(availability, 2),
            'total_customer_base': total_customers
        }
    
    def _calculate_comparisons(self, metrics):
        """Calculate meaningful comparisons based on deployment timeline and available data"""
        comparisons = {}
        
        # Deployment period performance vs 2024 baseline
        if 'since_deployment' in metrics and 'august_october_2024' in metrics:
            current = metrics['since_deployment']
            baseline_2024 = metrics['august_october_2024']
            
            comparisons['deployment_vs_2024_baseline'] = {
                'description': 'Current deployment period vs same period 2024',
                'supply_hours': {
                    'current': current['supply_hours']['average'],
                    'baseline_2024': baseline_2024['supply_hours']['average'],
                    'improvement_percent': self.calculate_delta(
                        current['supply_hours']['average'],
                        baseline_2024['supply_hours']['average']
                    )
                },
                'peak_load': {
                    'current': current['load_metrics']['peak_load_mw'],
                    'baseline_2024': baseline_2024['load_metrics']['peak_load_mw'],
                    'change_percent': self.calculate_delta(
                        current['load_metrics']['peak_load_mw'],
                        baseline_2024['load_metrics']['peak_load_mw']
                    )
                },
                'interruption_frequency': {
                    'current_daily_avg': current['interruptions']['total_interruptions'] / current['period']['days'],
                    'baseline_daily_avg': baseline_2024['interruptions']['total_interruptions'] / baseline_2024['period']['days'],
                    'change_percent': self.calculate_delta(
                        current['interruptions']['total_interruptions'] / current['period']['days'],
                        baseline_2024['interruptions']['total_interruptions'] / baseline_2024['period']['days']
                    )
                },
                'availability': {
                    'current': current['reliability']['availability_percentage'],
                    'baseline_2024': baseline_2024['reliability']['availability_percentage'],
                    'improvement_percent': self.calculate_delta(
                        current['reliability']['availability_percentage'],
                        baseline_2024['reliability']['availability_percentage']
                    )
                }
            }
        
        # Month-over-month comparison for 2025
        if 'september_2025' in metrics and 'october_2025' in metrics:
            sep_2025 = metrics['september_2025']
            oct_2025 = metrics['october_2025']
            
            comparisons['august_to_september_2025'] = {
                'description': 'September 2025 vs October 2025 (partial)',
                'supply_hours': {
                    'august': sep_2025['supply_hours']['average'],
                    'september': oct_2025['supply_hours']['average'],
                    'change_percent': self.calculate_delta(
                        oct_2025['supply_hours']['average'],
                        sep_2025['supply_hours']['average']
                    )
                },
                'interruptions': {
                    'august_daily_avg': sep_2025['interruptions']['total_interruptions'] / sep_2025['period']['days'],
                    'september_daily_avg': oct_2025['interruptions']['total_interruptions'] / oct_2025['period']['days'],
                    'change_percent': self.calculate_delta(
                        oct_2025['interruptions']['total_interruptions'] / oct_2025['period']['days'],
                        sep_2025['interruptions']['total_interruptions'] / sep_2025['period']['days']
                    )
                }
            }
        
        # Weekly trend analysis if we have enough data
        if 'current_week' in metrics and 'previous_week' in metrics:
            current_week = metrics['current_week']
            previous_week = metrics['previous_week']
            
            comparisons['week_over_week'] = {
                'description': 'Current week vs previous week',
                'supply_hours': {
                    'current': current_week['supply_hours']['average'],
                    'previous': previous_week['supply_hours']['average'],
                    'change_percent': self.calculate_delta(
                        current_week['supply_hours']['average'],
                        previous_week['supply_hours']['average']
                    )
                },
                'interruptions': {
                    'current': current_week['interruptions']['total_interruptions'],
                    'previous': previous_week['interruptions']['total_interruptions'],
                    'change_percent': self.calculate_delta(
                        current_week['interruptions']['total_interruptions'],
                        previous_week['interruptions']['total_interruptions']
                    )
                }
            }
        
        return comparisons
    
    def get_operational_insights(self, metrics):
        """Generate operational insights and recommendations specific to deployment period"""
        if self.stdout:
            self.stdout.write("Generating operational insights...")
        
        insights = {
            'deployment_performance': [],
            'performance_highlights': [],
            'areas_of_concern': [],
            'recommendations': [],
            'trends': [],
            'comparative_analysis': []
        }
        
        deployment_metrics = metrics['metrics_by_period'].get('since_deployment', {})
        october_metrics = metrics['metrics_by_period'].get('october_2025', {})
        september_metrics = metrics['metrics_by_period'].get('september_2025', {})
        comparisons = metrics.get('key_comparisons', {})
        
        # Deployment-specific insights
        if deployment_metrics:
            deployment_days = deployment_metrics.get('period', {}).get('days', 0)
            insights['deployment_performance'].append(
                f"Raven system operational for {deployment_days} days since August 9, 2025"
            )
            
            avg_supply = deployment_metrics.get('supply_hours', {}).get('average', 0)
            availability = deployment_metrics.get('reliability', {}).get('availability_percentage', 0)
            total_interruptions = deployment_metrics.get('interruptions', {}).get('total_interruptions', 0)
            
            insights['deployment_performance'].append(
                f"Overall deployment period average: {avg_supply:.1f} hours daily supply ({availability:.1f}% availability)"
            )
            insights['deployment_performance'].append(
                f"Total interruptions recorded: {total_interruptions} ({total_interruptions/deployment_days:.1f} per day average)"
            )
        
        # Performance analysis for October
        current_month_metrics = october_metrics if october_metrics else september_metrics
        if current_month_metrics:
            supply = current_month_metrics.get('supply_hours', {}).get('average', 0)
            availability = current_month_metrics.get('reliability', {}).get('availability_percentage', 0)
            month_name = "October 2025" if october_metrics else "September 2025"
            
            if supply >= 20:
                insights['performance_highlights'].append(
                    f"{month_name}: Excellent supply reliability at {supply:.1f} hours/day"
                )
            elif supply >= 16:
                insights['performance_highlights'].append(
                    f"{month_name}: Good supply performance at {supply:.1f} hours/day"
                )
            elif supply >= 12:
                insights['areas_of_concern'].append(
                    f"{month_name}: Moderate supply hours at {supply:.1f} hours/day - improvement needed"
                )
            else:
                insights['areas_of_concern'].append(
                    f"{month_name}: Low supply hours at {supply:.1f} hours/day - requires urgent attention"
                )
            
            # Interruption analysis
            interruptions = current_month_metrics.get('interruptions', {})
            fault_count = interruptions.get('fault_count', 0)
            load_shedding_count = interruptions.get('load_shedding_count', 0)
            
            if fault_count > 0 and load_shedding_count > 0:
                fault_ratio = fault_count / (fault_count + load_shedding_count) * 100
                if fault_ratio > 60:
                    insights['areas_of_concern'].append(
                        f"High fault-related interruptions: {fault_ratio:.1f}% of total interruptions"
                    )
                    insights['recommendations'].append(
                        "Implement proactive maintenance program to reduce equipment faults"
                    )
        
        # Comparative analysis with 2024 baseline
        baseline_comparison = comparisons.get('deployment_vs_2024_baseline', {})
        if baseline_comparison:
            supply_improvement = baseline_comparison.get('supply_hours', {}).get('improvement_percent', 0)
            availability_improvement = baseline_comparison.get('availability', {}).get('improvement_percent', 0)
            
            if supply_improvement > 10:
                insights['comparative_analysis'].append(
                    f"Supply hours improved by {supply_improvement:.1f}% compared to same period in 2024"
                )
                insights['performance_highlights'].append(
                    "Significant improvement over 2024 baseline performance"
                )
            elif supply_improvement > 0:
                insights['comparative_analysis'].append(
                    f"Modest supply hours improvement of {supply_improvement:.1f}% vs 2024"
                )
            elif supply_improvement < -10:
                insights['areas_of_concern'].append(
                    f"Supply hours declined by {abs(supply_improvement):.1f}% compared to 2024"
                )
            
            if availability_improvement > 5:
                insights['performance_highlights'].append(
                    f"System availability improved by {availability_improvement:.1f}% vs 2024 baseline"
                )
        
        # Month-to-month trend analysis
        monthly_comparison = comparisons.get('august_to_september_2025', {})
        if monthly_comparison:
            supply_change = monthly_comparison.get('supply_hours', {}).get('change_percent', 0)
            if abs(supply_change) > 5:
                trend_direction = "improved" if supply_change > 0 else "declined"
                insights['trends'].append(
                    f"Supply hours {trend_direction} by {abs(supply_change):.1f}% from previous month"
                )
        
        # Weekly trend analysis
        weekly_comparison = comparisons.get('week_over_week', {})
        if weekly_comparison:
            weekly_supply_change = weekly_comparison.get('supply_hours', {}).get('change_percent', 0)
            if abs(weekly_supply_change) > 10:
                trend = "improving" if weekly_supply_change > 0 else "declining"
                insights['trends'].append(
                    f"Recent weekly trend: {trend} ({weekly_supply_change:+.1f}%)"
                )
        
        # Generate specific recommendations
        if deployment_metrics:
            avg_interruption_duration = deployment_metrics.get('interruptions', {}).get('avg_duration_hours', 0)
            if avg_interruption_duration > 2:
                insights['recommendations'].append(
                    f"Average interruption duration of {avg_interruption_duration:.1f}h suggests need for faster restoration procedures"
                )
            
            peak_load = deployment_metrics.get('load_metrics', {}).get('peak_load_mw', 0)
            avg_load = deployment_metrics.get('load_metrics', {}).get('average_load_mw', 0)
            if peak_load > 0 and avg_load > 0:
                load_factor = (avg_load / peak_load) * 100
                if load_factor < 60:
                    insights['recommendations'].append(
                        f"Load factor of {load_factor:.1f}% indicates potential for demand management optimization"
                    )
        
        # Data quality and system insights
        insights['recommendations'].extend([
            "Continue monitoring system performance as deployment matures",
            "Establish benchmarks using 2024 historical data for ongoing comparisons",
            "Consider expanding Raven deployment to additional substations based on CLUB performance"
        ])
        
        return insights
    
    def generate_executive_summary(self, infrastructure, performance, insights):
        """Generate executive summary tailored to KEDCO deployment timeline"""
        deployment_metrics = performance['metrics_by_period'].get('since_deployment', {})
        october_metrics = performance['metrics_by_period'].get('october_2025', {})
        september_metrics = performance['metrics_by_period'].get('september_2025', {})
        
        # Use October data if available, otherwise September, otherwise deployment period
        current_metrics = october_metrics if october_metrics else (september_metrics if september_metrics else deployment_metrics)
        
        period_label = 'October 2025' if october_metrics else ('September 2025' if september_metrics else 'Since Deployment')
        
        return {
            'report_date': self.current_date.isoformat(),
            'substation': self.substation_name,
            'deployment_info': {
                'status': 'Progressive deployment - Technical module active',
                'deployment_date': self.deployment_start.isoformat(),
                'days_operational': self.days_since_deployment,
                'data_coverage': 'August 9, 2025 to October 6, 2025',
                'historical_baseline': 'Complete 2024 data available for comparison'
            },
            'infrastructure_summary': {
                'feeders': infrastructure['total_feeders'],
                'transformers': infrastructure['total_transformers'],
                'customers_served': infrastructure['total_customers'],
                'voltage_breakdown': infrastructure['feeder_breakdown']
            },
            'key_performance_indicators': {
                'period': period_label,
                'average_supply_hours': current_metrics.get('supply_hours', {}).get('average', 0),
                'system_availability': current_metrics.get('reliability', {}).get('availability_percentage', 0),
                'total_interruptions': current_metrics.get('interruptions', {}).get('total_interruptions', 0),
                'daily_interruption_rate': (current_metrics.get('interruptions', {}).get('total_interruptions', 0) / 
                                           current_metrics.get('period', {}).get('days', 1)),
                'peak_load_mw': current_metrics.get('load_metrics', {}).get('peak_load_mw', 0),
                'energy_delivered_mwh': current_metrics.get('energy_delivered', {}).get('total_mwh', 0)
            },
            'deployment_performance': insights['deployment_performance'][:3],
            'highlights': insights['performance_highlights'][:3],
            'concerns': insights['areas_of_concern'][:2],
            'recommendations': insights['recommendations'][:3],
            'comparative_insights': insights['comparative_analysis'][:2]
        }
    
    def generate_full_report(self):
        """Generate the complete management report"""
        if self.stdout:
            self.stdout.write(f"Generating full management report for {self.substation_name} substation...")
            self.stdout.write("=" * 60)
        
        try:
            # Get all report components
            infrastructure = self.get_infrastructure_overview()
            performance = self.get_performance_metrics()
            insights = self.get_operational_insights(performance)
            executive_summary = self.generate_executive_summary(infrastructure, performance, insights)
            
            # Compile full report
            full_report = {
                'executive_summary': executive_summary,
                'infrastructure_overview': infrastructure,
                'performance_metrics': performance,
                'operational_insights': insights,
                'metadata': {
                    'generated_at': datetime.now().isoformat(),
                    'generator_version': '1.0',
                    'substation_id': self.substation_id,
                    'data_sources': [
                        'HourlyLoad', 'FeederInterruption', 'DailyHoursOfSupply',
                        'FeederEnergyDaily', 'FeederEnergyMonthly'
                    ]
                }
            }
            
            if self.stdout:
                self.stdout.write("Report generation completed successfully!")
                self.stdout.write("=" * 60)
            
            return full_report
            
        except Exception as e:
            if self.stdout:
                self.stdout.write(f"Error generating report: {str(e)}")
            raise
